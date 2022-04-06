from pathlib import Path

import spacy
from docx import Document

nlp = spacy.load("en_core_web_lg")


def extract(document):
    for para in document.paragraphs:
        yield para.text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.text.replace("\n", " ")


if __name__ == "__main__":
    FILE_PATH = "../pretrained_data/2022_documents/raw"
    docs = Path(FILE_PATH).glob("*.docx")
    output = Path(FILE_PATH) / "raw_data.txt"
    handler = open(output, "w")
    count = 0
    for doc in docs:
        file = Document(doc.resolve())
        for text in extract(file):
            res = nlp(text)
            for sent in res.sents:
                if len(sent) <= 5:
                    pass
                else:
                    count += 1
                    print(sent, file=handler)
    handler.close()
    print(count)
